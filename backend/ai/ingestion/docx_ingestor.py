# docx_ingestor.py

from docx import Document as DocxDocument
from pathlib import Path

from ai.models.document import Document
from ai.models.source_type import SourceType


def ingest_docx(
    path: str,
    workspace_id: str
) -> Document:

    doc = DocxDocument(path)

    text = "\n".join(
        para.text
        for para in doc.paragraphs
    )

    return Document(
        workspace_id=workspace_id,
        source_type=SourceType.DOCX,
        title=Path(path).name,
        content=text,
        metadata={
            "paragraph_count": len(doc.paragraphs)
        }
    )